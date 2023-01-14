import os
import docx

#TODO: Format fonts, rounding, headings
#TODO: Load formatted Word template
#TODO: Add alternative export to csv or xlsx
#TODO: Subscript in table headings
#TODO: UserWarning: style lookup by style_id is deprecated. Use style name as key instead. return self._get_style_id_from_style(self[style_name], style_type)


class DataExport:
    def __init__(self):
        # Initialise the Word file
        self.doc = docx.Document()

    def add_table(self, data, heading=None, decimals=False, dba_alignment="left"):
        # Remove decimal points and 0
        if not decimals:
            data = data.astype(int)
        #TODO: Move A-weighted columns to right
        # Add the table heading
        assert heading is not None
        self.doc.add_heading(heading, 1)
        # Initialise the table
        table = self.doc.add_table(rows=(data.shape[0] + 1), cols=data.shape[1] + 1, style="Table Grid")
        # table.style.TableGrid   # Add in borders
        # Add dates in first column
        table.cell(0, 0).text = "Date"
        dates = data.index.tolist()
        for i in range(data.shape[0]):
            table.cell(i + 1, 0).text = str(dates[i])
        # Add column headings
        for j in range(data.shape[1]):
            heading = str(data.columns[j])  # Remove index params from spectral column headings
            heading = heading.split("Hz")
            if len(heading) > 1:
                heading = heading[0] + "Hz"
            else:
                heading = heading[0]
            table.cell(0, j + 1).text = heading
        # Loop over the DataFrame and assign data to the Word Table
        for i in range(data.shape[0]):    # For each row
            for j in range(data.shape[1]): # Go through each column
                # And assign the values in the table.
                cell = data.iat[i, j]
                table.cell(i + 1, j + 1).text = str(cell)

    def export(self, path=None, filename="results.docx"):
        assert path is not None
        path = os.path.join(path, filename)
        self.doc.save(path)
